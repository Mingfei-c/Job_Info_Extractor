"""
Resume Parser Service
Parse PDF and DOCX format resumes to extract text content
"""

import re
import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ResumeData:
    """Resume data structure"""
    raw_text: str           # Raw text
    file_path: str          # File path
    file_type: str          # pdf / docx
    char_count: int = 0     # Character count
    word_count: int = 0     # Word count
    
    def __post_init__(self):
        """Calculate statistics"""
        self.char_count = len(self.raw_text)
        self.word_count = len(self.raw_text.split())


class ResumeParser:
    """
    Resume Parser
    
    Supported formats:
    - PDF (.pdf)
    - DOCX (.docx)
    
    Usage:
        resume = ResumeParser.parse("my_resume.pdf")
        print(resume.raw_text)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}
    
    @classmethod
    def parse(cls, file_path: str) -> ResumeData:
        """
        Parse resume file
        
        Args:
            file_path: Resume file path (PDF or DOCX)
            
        Returns:
            ResumeData: Parsed resume data
            
        Raises:
            FileNotFoundError: File does not exist
            ValueError: Unsupported file format
        """
        path = Path(file_path)
        
        # Check if file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file format
        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
        
        # Parse based on format
        if ext == '.pdf':
            text = cls._parse_pdf(file_path)
            file_type = 'pdf'
        else:  # .docx
            text = cls._parse_docx(file_path)
            file_type = 'docx'
        
        # Clean text
        text = cls._clean_text(text)
        
        logger.info(f"Successfully parsed resume: {path.name} ({len(text)} characters)")
        
        return ResumeData(
            raw_text=text,
            file_path=str(path.absolute()),
            file_type=file_type
        )
    
    @staticmethod
    def _parse_pdf(path: str) -> str:
        """
        Parse PDF file
        
        Uses PyMuPDF (fitz) to extract text
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF is required to parse PDF files. "
                "Please run: pip install pymupdf"
            )
        
        text_parts = []
        
        with fitz.open(path) as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"PDF page {page_num}: {len(page_text)} characters")
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _parse_docx(path: str) -> str:
        """
        Parse DOCX file
        
        Uses python-docx to extract text (including paragraphs and tables)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required to parse DOCX files. "
                "Please run: pip install python-docx"
            )
        
        doc = Document(path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables (commonly used in resumes to display information)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean text
        
        - Remove excess whitespace
        - Normalize line endings
        - Remove special control characters
        """
        if not text:
            return ""
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove control characters (keep newlines and spaces)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Remove excess blank lines (more than 2 consecutive blank lines become 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove trailing whitespace from lines
        text = '\n'.join(line.rstrip() for line in text.split('\n'))
        
        return text.strip()


# ==================== Test Code ====================

if __name__ == "__main__":
    import sys
    
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    
    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <resume_file_path>")
        print("Supported formats: .pdf, .docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        resume = ResumeParser.parse(file_path)
        
        print("=" * 60)
        print(f"File: {resume.file_path}")
        print(f"Format: {resume.file_type}")
        print(f"Characters: {resume.char_count}")
        print(f"Words: {resume.word_count}")
        print("=" * 60)
        print("\n【Resume Content Preview (first 2000 characters)】\n")
        print(resume.raw_text[:2000])
        
        if len(resume.raw_text) > 2000:
            print(f"\n... ({len(resume.raw_text) - 2000} more characters)")
        
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except ImportError as e:
        print(f"Missing dependency: {e}")
        sys.exit(1)
